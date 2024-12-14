import os
import requests
import re
from inspect import stack
from urllib.parse import quote
from time import sleep
from datetime import datetime
from dotenv import load_dotenv
from bs4 import BeautifulSoup
from typing import Literal
from colorama import Fore, Style, init
import asyncio

if __name__ == '__main__' or '.' not in __name__:
    from log import log
else:
    from .log import log

load_dotenv(os.path.join(os.path.dirname(__file__), '.env'))
path = '\\'.join(os.path.dirname(os.path.abspath(__file__)).split('\\')[:-1])



# =========================< FILE TO TEXT >=========================
# pdf, docx, mp3 to text. TODO: summary in large document (or embbeding)
import PyPDF2
import docx
from groq import Groq

groq_api_key = os.getenv('GROQ_API_KEY')
groq_client = Groq(api_key=groq_api_key)


def speech_recognition(file_name: str) -> str:  # TODO: local_whisper

    with open(file_name, "rb") as file:
        translation = groq_client.audio.transcriptions.create(
        file=(file_name, file.read()),
        model="whisper-large-v3")
        
        text = translation.text
    
    return str(text).strip()


def pdf_to_text(pdf_path: str) -> str:
    with open(pdf_path, 'rb') as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        extracted_text = ""
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text:
                extracted_text += text

    return extracted_text


def docx_to_text(docx_path: str) -> str:
    doc = docx.Document(docx_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)

    return '\n'.join(full_text)


def files_to_text(files: list | str) -> str:
    'Turns text files or audio into text'
    if type(files) == str:
        files = [files]

    result = ''
    
    try:
        for file_path in files:

            if file_path.endswith('.pdf'):
                text = pdf_to_text(file_path)
            
            elif file_path.endswith('.docx'):
                text = docx_to_text(file_path)
            
            elif file_path.endswith('.mp3'):
                text = speech_recognition(file_path)

            else:
                with open(file_path, 'r') as file:
                    text = file.read()
            
            result += f'{file_path}:\n{text}\n\n'
            
        return result
    
    except Exception as e:
        log(e, error=True)
        return f'Error reading {file_path}'



# =========================< CONVERT FILES >=========================
import fitz
def pdf_to_image(pdf_path: str, quality: int = 3, extension: str = 'png') -> list[str]:
    file_name = pdf_path[:-4]
    doc = fitz.open(pdf_path)
    photos = []
    count = len(doc)
        
    for i in range(count):
        page = doc.load_page(i)
        pix = page.get_pixmap(matrix=fitz.Matrix(quality, quality))
        temp_file_name = f"{file_name}_{i}.{extension}"
        pix.save(temp_file_name)
        photos.append(temp_file_name)

    doc.close()
    log(f'Success({count} pages): {photos}')
    return photos


def merges_pdf(files: list[str]) -> str:
    merger = fitz.Document()

    for file in files:
        doc = fitz.open(file)
        merger.insert_pdf(doc)
        doc.close()

    merged_file = f'{files[0][:-4]}_merged.pdf'
    merger.save(merged_file)

    log(f'Success({len(files)} files): {merged_file}')
    return merged_file




# =========================< GROQ API >=========================
def groq_api(messages: list, model: str = 'llama-3.1-70b-versatile') -> str:
    # https://console.groq.com/docs/models
    response = groq_client.chat.completions.create(
        messages=messages,
        model=model)
        
    return str(response.choices[0].message.content)



# =========================< GOOGLE API >=========================
import google.generativeai as genai
import PIL.Image

os.environ["GOOGLE_API_KEY"] = str(os.getenv("GOOGLE_API_KEY"))
genai.configure(api_key=os.environ['GOOGLE_API_KEY'])
model = genai.GenerativeModel("gemini-1.5-flash-002")
model = genai.GenerativeModel("gemini-2.0-flash-exp")


def genai_api(messages: list[dict] | str, file_paths: list | str = []) -> str:

    if type(messages) == str:
        user_message = messages
        formatted_history = []

    else:  # google is a bit special and to keep everything in the same style I do the transformation here.
        user_message = messages[-1]['content']
        history = messages[:-1]

        formatted_history = []
        for message in history:
            role = 'user' if message['role'] in ['user', 'system'] else 'model' 
            content = message['content']
            formatted_history.append({"role": role, "parts": content})


    if type(file_paths) == str:
        file_paths = [file_paths]

    
    chat = model.start_chat(history=formatted_history)

    files = []
    for file_path in file_paths:
        if file_path.endswith('.png') or file_path.endswith('.jpg') or file_path.endswith('.jpeg') or file_path.endswith('.webp'):
            files.append(PIL.Image.open(file_path))

        else:
            user_message += '\n\n' + file_path + ':\n' + files_to_text(file_path)

    if file_paths:
        response = chat.send_message([user_message] + files)

    else:
        response = chat.send_message(user_message)


    return response.text


# =========================< LLM API >=========================
def llm_api(messages: list[dict], files: str | list = [], provider: Literal['groq', 'google'] = 'groq'):

    if type(files) == str:
        files = [files]
        
    start_time = datetime.now()
    
    if provider == 'google' or files:
        try:
            answer = genai_api(messages, files)
        except Exception as e:
            log(f'google error: {e}', error=True)
            answer = groq_api(messages)
    
    else:
        answer = groq_api(messages)

    log(f'answer: {answer}, messages: {messages}, files: {files}, provider: {provider}, time: {datetime.now()-start_time}')

    return answer



# =========================< WOLFRAM ALPHA >=========================
import requests
import re

WOLFRAM_SIMPLE_API = os.getenv('WOLFRAM_SIMPLE_API_KEY')
WOLFRAM_SHOW_STEPS_RESULT = os.getenv('WOLFRAM_SHOW_STEPS_RESULT')

# TODO: remake prompt, add an interpreter
prompt_for_transform_query_wolfram = """You have to turn a user query into a query that wolfram alpha will understand (If the enquiry is normal, leave it that way).

Examples:

1. User: Use wolfram alpha and solve this equation 3x-1=11
   You: 3x-1=11

2. User: Solve 3x^2-7x+4=0
   You: Solve 3x^2-7x+4=0

3. User: {sqrt2}}^{sqrt{2}
   You: sqrt[2]^sqrt[2]

4. User: Slve 2x^2=16
   You: Solve 2x^2=16

5. User: Help me and do the math (567^34)/435-6758
   You: (567^34)/435-6758

6. User: Sum of roots 8x^3-4x^2+11x-36=0
   You: Sum of roots 8x^3-4x^2+11x-36=0

7. User: Use wolfram alpha to find the population of France
   You: France population


Also wolfram alpha only accepts the English language

8. User: x в кубе равно 8
   You: x^3=8

9. User: 14x meno 5 uguale 0
   You: 14x-5=0

10. User: Calculer 456^45
   You: Calculate 456^45


Now, respond to the query by following the rules
"""


def calculator(expression: str) -> str:
    try:
        expression = expression.replace('^', '**')
        log(expression)
        return str(eval(expression))
    
    except Exception as e:
        return f"Calculator error: {e}"
    
    
def download_image(url: str) -> str:
    'Downloads the image from the link. Returns the name of the downloaded image (name is time).'
    response = requests.get(url)
    file_name = f'{hash(url)}.png'
    with open(file_name, 'wb') as file:
        file.write(response.content)
    log(file_name)
    return file_name


def wolfram_short_answer_api(text: str) -> str:
    'https://products.wolframalpha.com/short-answers-api/documentation'
    query = quote(text)
    url = f'https://api.wolframalpha.com/v1/result?appid={WOLFRAM_SIMPLE_API}&i={query}'
    answer = requests.get(url).text
    log(answer)
    return answer


def wolfram_llm_api(text: str) -> tuple[str, list]:
    'https://products.wolframalpha.com/llm-api/documentation - text version of the WolframAlpha page answer. Returns the text and a list of links to images'
    query = quote(text)
    url = f'https://www.wolframalpha.com/api/v1/llm-api?input={query}&appid={WOLFRAM_SHOW_STEPS_RESULT}'
    answer = requests.get(url).text
    
    answer = answer[:answer.find('Wolfram|Alpha website result for "')]
    links = re.findall(r'https?://\S+', answer)
    answer = re.sub(r'https?://\S+', 'Images will be attached to the answer', answer)
    log(f'{answer}, {links}')
    return answer, links


def wolfram_simple_api(text: str) -> str:
    'https://products.wolframalpha.com/simple-api/documentation - WolframAlpha answer page image'
    query = quote(text)
    link = f'https://api.wolframalpha.com/v1/simple?appid={WOLFRAM_SIMPLE_API}&i={query}%3F'
    file_name = download_image(link)
    log(file_name)
    return file_name


def wolfram_short_answer(query: str) -> str:
    'Short answer from WolframAlpha'
    quick_answer = wolfram_short_answer_api(query)

    if 'Wolfram|Alpha did not understand' in quick_answer:
        log(f'Wolfram Alpha did not understand {query}. Ask llm', error=True)
        messages = [
            {'role': 'user', 'content': prompt_for_transform_query_wolfram},
            {'role': 'user', 'content': query}
        ]
        query = groq_api(messages=messages, model='llama3-8b-8192')
        quick_answer = wolfram_short_answer_api(query)

    log(quick_answer)
    return quick_answer


def wolfram_full_answer(text: str):  # TODO: async + wolfram_simple_api
    'returns the text version of the answer sheet, the image links and the answer sheet as a picture'

    full_answer, full_answer_images = wolfram_llm_api(text)

    if 'Alpha did not understand your input' in full_answer:
        log(f'Wolfram Alpha did not understand {text}. Ask llm', error=True)
        messages = [
            {'role': 'user', 'content': prompt_for_transform_query_wolfram},
            {'role': 'user', 'content': text}
        ]
        text = groq_api(messages=messages, model='llama3-8b-8192')
        
        full_answer, full_answer_images = wolfram_llm_api(text)
        
    images = full_answer_images # TODO: [wolfram_simple_api(text)] + full_answer_images 
    log(f'{full_answer}, {images}')
    return full_answer, images



# =========================< TAVILY AND DUCKDUCKGO (INTERNERT) >=========================
# internetfrom tavily import TavilyClient
from tavily import TavilyClient
from duckduckgo_search import DDGS
from deep_translator import GoogleTranslator, single_detection

detect_language_api_key = os.getenv('DETECT_LANGUAGE_API')
tavily_client = TavilyClient(api_key=os.getenv('TAVILY_API_KEY'))
# TODO: ggcs translator, detect language


def detect_language(text: str) -> str:  # TODO: google and ggcs translator
    try:
        return single_detection(text=text, api_key=detect_language_api_key, detailed=False)
    except:
        return 'en'
    

def parsing(links: str | list) -> str:
    start_time = datetime.now() 
    try:
        responce = tavily_client.extract(urls=links)
        return [r['raw_content'] for r in responce['results']]
    
    except Exception as e:
        if type(links) == str:
            links = [links]
        result = ''
        for link in links:
            try:
                responce = tavily_client.extract(urls=link)
                result += f'{link}: {responce["result"][0]["raw_content"]}\n'
            except:
                result += f'{link}: Error when extracting text ({e})\n'

    log(f'{links}, {result}, {datetime.now() - start_time}')
    return result
    

def DDGS_answer(text: str) -> str:
    'A short answer like Google. Sometimes nothing comes up. But mostly the answer comes from wikipedia.'
    try:
        response = DDGS().answers(text)[:3]
        return '\n'.join([r["text"] for r in response])
    except:
        return ''


def DDGS_images(text: str, max_results: int = 9) -> list[str]:
    images = [i['image'] for i in DDGS().images(text, max_results=max_results)]
    return images


def tavily_answer(query: str):
    response = tavily_client.qna_search(query=query)
    return response


def tavily_get_context(query: str, topic = 'news') -> str:
    return tavily_client.get_search_context(query=query, topic=topic).replace('\\', '')


def tavily_content(text: str, max_results: int = 4):
    response = tavily_client.search(query=text, search_depth='basic', max_results=max_results)['results']
    results = ''
    for r in response:
        results += f'{r["url"]}: {r["content"]}\n'

    return results


prompt_for_sum = f"""You are a precise summarization expert. Your task is to create a clear and relevant summary based on the provided text and query.

Context: I will provide you with:
1. Query: A specific question or topic of interest
2. Source text: Content from a webpage or document or a collection of several responses

Ignore unnecessary information and answer the query well.
"""


def sum_page(link: str, query: str) -> str:
    prompt = prompt_for_sum + f'\n\nQuery:\n{query}\n\nSource text:\n{parsing(link)}'  # TODO: parser error
    messages = [{"role": "user", "content": prompt}]
    try:
        llm_answer = llm_api(messages=messages, provider='google')
        log(llm_answer)
    except Exception as e:
        l = len(str(messages))
        llm_answer = 'Error'
        log(f'{e}. len: {l}', error=True)

    return f'{link}: {llm_answer}\n\n\n'


def google_links(text: str, max_results: int = 5) -> list[str]:
    try:
        links = [i['href'] for i in DDGS().text(text, max_results=max_results)]
        log(str(links))
        return links
    except Exception as e:
        log(f'{e}', error=True)
        return []


def google_short_answer(text: str) -> str:
    resp = DDGS_answer(text)
    final_answer = resp if resp else tavily_answer(text)
    log(final_answer)
    return final_answer
    

def google_full_answer(text: str, max_results: int = 5):  # TODO: async, for lyrics

    start_time = datetime.now()
    links = google_links(text=text, max_results=max_results)

    sum_answers = ''
    for link in links:
        sum_answers += sum_page(link=link, query=text)  # TODO: max_result excluding parsing error

    prompt = prompt_for_sum + f"""Summarize the text above in a concise and relevant way. Ensure that the summary is well-organized and captures the main points of the text. The summary should be based on the query and the provided text. If the text is not relevant to the query, please provide a summary that is not relevant to the query. Source text will be composed of several responses. Write the final text"""
    prompt +=  f'\n\nQuery:\n{text}\n\nSource text:\n{sum_answers}'
    
    final_messages = [{"role": "user", "content": prompt}]
    final_answer = llm_api(messages=final_messages, provider='groq')

    log(f'{final_answer}, {datetime.now() - start_time}')
    return final_answer


def google_news(text: str):
    news_content = tavily_get_context(text)
    
    prompt = prompt_for_sum + f"""Turn all these texts into one\n\nQuery:\n{text}\n\nSource text:\n{news_content}"""

    final_messages = [{"role": "user", "content": prompt}]
    final_answer = llm_api(messages=final_messages, provider='groq')

    log(final_answer)
    return final_answer



import aiohttp
import asyncio

async def download_image(session, url, save_path):
    try:
        async with session.get(url) as response:
            if response.status == 200:
                with open(save_path, 'wb') as file:
                    file.write(await response.read())
                return save_path
            else:
                print(f"Failed to download {url}: HTTP {response.status}")
    except Exception as e:
        print(f"Error downloading {url}: {e}")


async def download_images(image_urls, save_dir):
    os.makedirs(save_dir, exist_ok=True)
    
    async with aiohttp.ClientSession() as session:
        tasks = []
        for i, url in enumerate(image_urls):
            try:
                save_path = os.path.join(save_dir, f"image_{i + 1}.jpg")
                tasks.append(download_image(session, url, save_path))
            except:
                pass
        
        return await asyncio.gather(*tasks)


async def google_image(text, max_results=9, download_image=True):
    urls = DDGS_images(text, max_results=max_results)
    start_time = datetime.now()
    

    if not download_image:
        log(urls)
        return f'google_image: The {len(urls)} of images on the {text} query will be prefixed to your response', urls

    
    file_paths = await download_images(urls, 'images')
    
    log(f'{file_paths}, {datetime.now()-start_time}')
    return f'The {len(file_paths)} of the {text} query images will be appended to the response. Answer other user questions, tell information about the query, or say something like images found. DON\'T write anything like [Image of ...] or you\'ll be shut down.', file_paths



from youtube_transcript_api import YouTubeTranscriptApi



# =========================< YOUTUBE >=========================
def transcript2text(transcript: list[dict], sep='\n'):
    result = []
    current_minute = -1

    for t in transcript:
        minute = int(t["start"] // 60)
        if minute != current_minute:
            result.append(f"{minute} minute:")
            current_minute = minute
        
        result.append(t["text"])

    return sep.join(result)



def get_youtube_transcripts(link: str, language: str = 'en'):

    if 'youtube.com' in link:
        # https://www.youtube.com/watch?v=xxxxxxxxxxx -> xxxxxxxxxxx
        pattern_for_youtube_video = r"(?:v=|\/)([a-zA-Z0-9_-]{11})"  # unreadable, but it works perfectly
        match = re.search(pattern_for_youtube_video, link)
        if match:
            video_id = match.group(1)
        else:
            return 'Error: incorrect link'
    else:
        video_id = link

    
    # get title
    response = requests.get(link)
    soup = BeautifulSoup(response.text, 'html.parser')
    title_tag = soup.find("title")
    if title_tag:
        title = title_tag.text.replace(" - YouTube", "")
    else:
        title = 'Error: title not found'

    # get transcript: [{'text': str, 'start': float, 'duration': float}, ...]
    try:  
        transcript = YouTubeTranscriptApi.get_transcript(video_id, languages=[language])
        log(f'1 (defult language): {language}, {link}({title})')
        return transcript2text(transcript), title
    except:
        pass


    try:
        language = detect_language(title)
        transcript = YouTubeTranscriptApi.get_transcript(video_id, languages=[language])
        log(f'2 (detect language): {language}, {link}({title})')
        return transcript2text(transcript), title
    except:
        pass


    try:
        transcripts = YouTubeTranscriptApi.list_transcripts(video_id)
        for transcript in transcripts:
            transcript = transcript.translate('en').fetch()
            break
        log(f'3 (translate). {link}({title})')
    except Exception as e:
        log(f'error: {e}', error=True)
        transcript = [{'text': f'Subtitles are not available', 'start': 0.0}]
    
    return transcript2text(transcript), title


def youtube_sum(link: str, question: str | None = None, language: str = 'en') -> str: 
    # TODO: solve the problem with periodic errors (just does not give subtitles, and then gives them) and make it possible to ask questions (llm_select_tool).
    text, title = get_youtube_transcripts(link, language)

    if question:
        content = f'Answer the question "{question}" based on this YouTube video "{title}": {text}'
    else:
        content = f'Summarize this YouTube video "{title}": {text}'

    try:
        answer = llm_api(messages=[{'role': 'user', 'content': content}], provider='google')
        log(f'{link} ({question}): {answer}')
        return answer
    except Exception as e:
        log(f'{link} error ({e})', error=True)
        return f'Error({e}). Most like too much text or a completely incomprehensible error that cannot be fixed (or subtitles are not available). Tell the user to try again later'



